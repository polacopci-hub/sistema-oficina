import flet as ft
from supabase import create_client, Client
from fpdf import FPDF
from docx import Document 
from docx.shared import Pt
import datetime
from datetime import timedelta
import os
import re
import urllib.parse 

# --- CONFIGURAÇÃO DO BANCO ---
SUPABASE_URL = "https://mwqwceayaouowgehuukf.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Im13cXdjZWF5YW91b3dnZWh1dWtmIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NzA2NjYzMDEsImV4cCI6MjA4NjI0MjMwMX0.5ItH8uAEEcHxDkew18e_kaGFkgIkfp5LaMM60RjT0U0"

try:
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
except Exception as e:
    print(f"Erro no Banco: {e}")

def main(page: ft.Page):
    page.title = "Oficina App"
    page.theme_mode = ft.ThemeMode.LIGHT
    page.scroll = "adaptive"
    page.window_width = 390
    page.window_height = 844

    usuario_atual = {"id": None, "nome": None, "setor": None}
    dados_atuais = []
    id_em_edicao = {"id": None} 

    def limpar_nome_arquivo(texto):
        s = texto.replace(" ", "_")
        return re.sub(r'[^a-zA-Z0-9_\-]', '', s)

    # --- GERADOR DE PDF ---
    def gerar_pdf_nuvem(lista_dados, periodo, nome_usuario):
        try:
            pdf = FPDF() 
            pdf.add_page()
            pdf.set_font("helvetica", "B", 16)
            
            pdf.cell(0, 10, "RELATORIO DE SERVICOS", align="C", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("helvetica", "", 10)
            pdf.cell(0, 10, f"Periodo: {periodo}", align="C", new_x="LMARGIN", new_y="NEXT")
            
            titulo_tecnico = f"Tecnico: {nome_usuario}"
            if nome_usuario == "TODOS": titulo_tecnico = "Relatorio Geral (Todos os Tecnicos)"
            pdf.cell(0, 10, titulo_tecnico, align="C", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(5)

            pdf.set_fill_color(220, 220, 220)
            pdf.set_font("helvetica", "B", 8) 
            pdf.cell(20, 8, "DATA", border=1, fill=True, align="C")
            pdf.cell(20, 8, "PLACA", border=1, fill=True, align="C")
            pdf.cell(30, 8, "MODELO", border=1, fill=True, align="C")
            pdf.cell(35, 8, "CLIENTE", border=1, fill=True, align="C")
            pdf.cell(0, 8, "OBSERVACOES", border=1, fill=True, align="C", new_x="LMARGIN", new_y="NEXT")

            pdf.set_font("helvetica", "", 7)
            for item in lista_dados:
                data_fmt = f"{item['data_hora'][8:10]}/{item['data_hora'][5:7]}"
                modelo = str(item.get('modelo', '-'))[:15]
                cliente = str(item.get('cliente', '-'))[:20]
                obs_texto = (str(item.get('observacoes', '-'))[:55] + '..') if len(str(item.get('observacoes', '-'))) > 55 else str(item.get('observacoes', '-'))
                pdf.cell(20, 8, data_fmt, border=1, align="C")
                pdf.cell(20, 8, str(item.get('placa', '-')), border=1, align="C")
                pdf.cell(30, 8, modelo, border=1, align="C")
                pdf.cell(35, 8, cliente, border=1, align="C") 
                pdf.cell(0, 8, obs_texto, border=1, new_x="LMARGIN", new_y="NEXT")

            data_hoje = datetime.datetime.now().strftime("%Y-%m-%d")
            nome_arq = f"Relatorio_{limpar_nome_arquivo(nome_usuario)}_{data_hoje}.pdf"
            pdf.output(nome_arq)

            with open(nome_arq, "rb") as f:
                supabase.storage.from_("relatorios").upload(
                    path=nome_arq, file=f, file_options={"content-type": "application/pdf", "x-upsert": "true"}
                )
            url = supabase.storage.from_("relatorios").get_public_url(nome_arq)
            os.remove(nome_arq) 
            return url
        except Exception as ex:
            print(f"Erro PDF: {ex}")
            return None

    # --- GERADOR DE WORD (COMISSÃO - CORRIGIDO) ---
    def gerar_word_nuvem(lista_dados, periodo, nome_usuario):
        try:
            doc = Document()
            
            titulo = doc.add_paragraph()
            run = titulo.add_run(f"{periodo}\nCOMISSÃO do SERVIÇO REALIZADO – {nome_usuario}")
            run.bold = True
            run.font.size = Pt(14)
            
            # --- SUPER FILTRO CORRIGIDO ---
            normais = []
            extras = []
            
            for item in lista_dados:
                val_extra = item.get('is_extra')
                eh_extra = False
                
                # Verifica se é True booleano
                if isinstance(val_extra, bool) and val_extra is True:
                    eh_extra = True
                # Verifica se é texto "true" (caso o banco mande como string)
                elif str(val_extra).lower() in ['true', '1', 'yes', 't']:
                    eh_extra = True
                
                if eh_extra:
                    extras.append(item)
                else:
                    normais.append(item)

            print(f"DEBUG WORD: Total {len(lista_dados)} | Normais {len(normais)} | Extras {len(extras)}")

            # --- TABELA 1: SERVIÇO NORMAL ---
            tabela = doc.add_table(rows=1, cols=5)
            tabela.style = 'Table Grid'
            hdr_cells = tabela.rows[0].cells
            hdr_cells[0].text = 'Veículo'
            hdr_cells[1].text = 'Placa'
            hdr_cells[2].text = 'Serviço'
            hdr_cells[3].text = 'Proprietário'
            hdr_cells[4].text = 'R$'

            for item in normais:
                row_cells = tabela.add_row().cells
                row_cells[0].text = str(item.get('modelo', ''))
                row_cells[1].text = str(item.get('placa', ''))
                row_cells[2].text = str(item.get('observacoes', ''))
                row_cells[3].text = str(item.get('cliente', ''))
                row_cells[4].text = '' 
            
            row_total = tabela.add_row().cells
            row_total[3].text = "TOTAL"
            row_total[4].text = ""

            doc.add_paragraph("\n")

            # --- TABELA 2: EXTRA ---
            titulo_extra = doc.add_paragraph()
            run_ex = titulo_extra.add_run("Extra")
            run_ex.bold = True
            
            tabela_ex = doc.add_table(rows=1, cols=5)
            tabela_ex.style = 'Table Grid'
            hdr_ex = tabela_ex.rows[0].cells
            hdr_ex[0].text = 'Veículo'
            hdr_ex[1].text = 'Placa'
            hdr_ex[2].text = 'Serviço'
            hdr_ex[3].text = 'Proprietário'
            hdr_ex[4].text = 'R$'

            for item in extras:
                row_cells = tabela_ex.add_row().cells
                row_cells[0].text = str(item.get('modelo', ''))
                row_cells[1].text = str(item.get('placa', ''))
                row_cells[2].text = str(item.get('observacoes', ''))
                row_cells[3].text = str(item.get('cliente', ''))
                row_cells[4].text = ''
            
            row_total_ex = tabela_ex.add_row().cells
            row_total_ex[3].text = "TOTAL"
            row_total_ex[4].text = ""

            data_hoje = datetime.datetime.now().strftime("%Y-%m-%d")
            nome_arq = f"Comissao_{limpar_nome_arquivo(nome_usuario)}_{data_hoje}.docx"
            doc.save(nome_arq)

            with open(nome_arq, "rb") as f:
                supabase.storage.from_("relatorios").upload(
                    path=nome_arq, file=f, file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "x-upsert": "true"}
                )
            url = supabase.storage.from_("relatorios").get_public_url(nome_arq)
            os.remove(nome_arq)
            return url

        except Exception as ex:
            print(f"Erro Word: {ex}")
            return None

    # --- TELA DE CADASTRO ---
    def tela_cadastro():
        page.clean()
        txt_novo_nome = ft.TextField(label="Seu Nome Completo (Login)")
        txt_novo_setor = ft.TextField(label="Seu Setor (Ex: Pintura)")
        txt_novo_pin = ft.TextField(label="Crie uma Senha PIN (Números)", keyboard_type="number", password=True)
        txt_msg_erro = ft.Text("", color="red") 
        btn_salvar_cad = ft.FilledButton("SALVAR CADASTRO", width=300, height=50)

        def salvar_usuario(e):
            if not txt_novo_nome.value or not txt_novo_pin.value:
                txt_msg_erro.value = "Erro: Preencha os campos!"; page.update(); return
            try:
                nome_limpo_cad = txt_novo_nome.value.strip()
                pin_limpo_cad = txt_novo_pin.value.strip()
                supabase.table("usuarios").insert({
                    "nome": nome_limpo_cad, 
                    "pin": pin_limpo_cad, 
                    "setor": txt_novo_setor.value, 
                    "ativo": True
                }).execute()
                tela_login()
            except: txt_msg_erro.value = "Erro ao cadastrar"; page.update()

        btn_salvar_cad.on_click = salvar_usuario
        page.add(ft.Column([
            ft.Text("CRIAR CONTA", size=25, weight="bold", color="blue"),
            ft.Text("Preencha seus dados", size=14, color="grey"),
            ft.Divider(color="transparent"),
            txt_novo_nome, txt_novo_setor, txt_novo_pin,
            txt_msg_erro, 
            ft.Divider(color="transparent"),
            btn_salvar_cad,
            ft.TextButton("Voltar para Login", on_click=lambda e: tela_login())
        ], alignment="center", horizontal_alignment="center"))

    # --- TELA DE LOGIN ---
    def tela_login():
        page.clean()
        txt_login_nome = ft.TextField(label="Digite seu Nome de Usuário", width=300)
        txt_login_pin = ft.TextField(label="Senha PIN", password=True, width=150, text_align="center", keyboard_type="number")
        txt_aviso_login = ft.Text("", color="red", size=16, weight="bold")
        btn_entrar = ft.FilledButton("ENTRAR", width=200, height=50)

        def logar(e):
            txt_aviso_login.value = "Verificando..."; page.update()
            nome_limpo = txt_login_nome.value.strip() 
            pin_limpo = txt_login_pin.value.strip()   
            
            if not nome_limpo or not pin_limpo:
                txt_aviso_login.value = "Preencha os campos!"; page.update(); return

            res = supabase.table("usuarios").select("*").ilike("nome", f"{nome_limpo}%").execute()
            
            usuario_encontrado = None
            if res.data:
                for u in res.data:
                    if u['pin'] == pin_limpo:
                        usuario_encontrado = u
                        break
            
            if usuario_encontrado:
                usuario_atual.update(usuario_encontrado)
                sistema_principal()
            else:
                txt_aviso_login.value = "Senha ou Usuário Incorretos!"; page.update()

        btn_entrar.on_click = logar
        page.add(ft.Column([
            ft.Text("OFICINA APP", size=30, weight="bold", color="blue"),
            ft.Text("Login Seguro", size=12, color="grey"),
            ft.Divider(height=20, color="transparent"),
            txt_login_nome,
            txt_login_pin,
            txt_aviso_login,
            btn_entrar,
            ft.Divider(),
            ft.Text("Não tem acesso?"),
            ft.OutlinedButton("CRIAR CONTA NOVA", on_click=lambda e: tela_cadastro(), width=200)
        ], alignment="center", horizontal_alignment="center"))

    # --- SISTEMA PRINCIPAL ---
    def sistema_principal():
        page.clean()
        
        lbl_titulo_os = ft.Text("NOVA ORDEM DE SERVICO", size=18, weight="bold")
        txt_placa = ft.TextField(label="Placa")
        txt_modelo = ft.TextField(label="Modelo")
        txt_cliente = ft.TextField(label="Cliente")
        txt_obs = ft.TextField(label="O que foi feito? (Observações)", multiline=True, min_lines=3, max_lines=5)
        
        chk_extra = ft.Checkbox(label="Serviço Extra / Pós-Horário", value=False)
        
        btn_salvar_os = ft.FilledButton("SALVAR REGISTRO", width=300, height=50)
        btn_cancelar_edicao = ft.TextButton("Cancelar Edição", visible=False)

        def resetar_form():
            id_em_edicao['id'] = None; lbl_titulo_os.value = "NOVA ORDEM DE SERVICO"; lbl_titulo_os.color = "black"
            btn_salvar_os.text = "SALVAR REGISTRO"; btn_cancelar_edicao.visible = False
            txt_placa.value = ""; txt_modelo.value = ""; txt_cliente.value = ""; txt_obs.value = ""
            chk_extra.value = False 
            page.update()

        def salvar_os(e):
            if not txt_obs.value or not txt_placa.value: return
            dados = {
                "usuario_id": usuario_atual['id'], 
                "placa": txt_placa.value, 
                "modelo": txt_modelo.value, 
                "cliente": txt_cliente.value, 
                "observacoes": txt_obs.value,
                "is_extra": chk_extra.value 
            }
            if id_em_edicao['id']:
                supabase.table("servicos").update(dados).eq("id", id_em_edicao['id']).execute()
            else:
                supabase.table("servicos").insert(dados).execute()
            resetar_form(); page.snack_bar = ft.SnackBar(ft.Text("Registro Salvo!")); page.snack_bar.open = True; page.update()

        btn_salvar_os.on_click = salvar_os
        btn_cancelar_edicao.on_click = lambda e: resetar_form()
        container_nova_os = ft.Column([lbl_titulo_os, txt_placa, txt_modelo, txt_cliente, ft.Divider(), txt_obs, chk_extra, btn_salvar_os, btn_cancelar_edicao])

        txt_dt_ini = ft.TextField(label="Início", value=str(datetime.date.today()), width=140)
        txt_dt_fim = ft.TextField(label="Fim", value=str(datetime.date.today()), width=140)
        dd_filtro_func = ft.Dropdown(label="Filtrar por Técnico", width=300, visible=False)
        lista_cards = ft.Column()
        
        btn_gerar = ft.FilledButton("GERAR RELATÓRIO", visible=False, width=300)
        txt_feedback_pdf = ft.Text("", color="blue")
        linha_botoes_pdf = ft.Row(visible=False, alignment=ft.MainAxisAlignment.CENTER, wrap=True) 

        def buscar(e):
            lista_cards.controls.clear(); dados_atuais.clear(); 
            btn_gerar.visible = False
            linha_botoes_pdf.visible = False 
            txt_feedback_pdf.value = ""
            page.update()

            q = supabase.table("servicos").select("*").gte("data_hora", f"{txt_dt_ini.value}T00:00:00").lte("data_hora", f"{txt_dt_fim.value}T23:59:59").order("id", desc=True)
            if dd_filtro_func.visible and dd_filtro_func.value != "todos" and dd_filtro_func.value:
                q = q.eq("usuario_id", dd_filtro_func.value)
            elif not dd_filtro_func.visible:
                q = q.eq("usuario_id", usuario_atual['id'])
            
            res = q.execute()
            if res.data:
                dados_atuais.extend(res.data); btn_gerar.visible = True
                for item in res.data:
                    texto_extra = " [EXTRA]" if item.get('is_extra') else ""
                    card = ft.Container(
                        padding=10, border=ft.Border.all(1, "grey"), border_radius=8,
                        content=ft.Column([
                            ft.Row([ft.Text(f"PLACA: {item['placa']}{texto_extra}", weight="bold"), ft.Text(f"{item['data_hora'][8:10]}/{item['data_hora'][5:7]}", color="grey")], alignment="spaceBetween"),
                            ft.Text(f"Veículo: {item.get('modelo','-')}"),
                            ft.Text(f"Cliente: {item.get('cliente','-')}", size=12),
                            ft.Text(f"Feito: {item.get('observacoes','-')}", color="blue"),
                            ft.Divider(),
                            ft.Row([
                                ft.TextButton("[EDITAR]", on_click=lambda e, i=item: preparar_edicao(i)),
                                ft.TextButton("[EXCLUIR]", on_click=lambda e, idx=item['id']: (supabase.table("servicos").delete().eq("id", idx).execute(), buscar(None)), style=ft.ButtonStyle(color="red"))
                            ], alignment="end")
                        ])
                    )
                    lista_cards.controls.append(card)
            page.update()

        def preparar_edicao(item):
            id_em_edicao['id'] = item['id']; txt_placa.value = item['placa']; txt_modelo.value = item['modelo']
            txt_cliente.value = item['cliente']; txt_obs.value = item['observacoes']
            chk_extra.value = item.get('is_extra', False)
            
            lbl_titulo_os.value = f"EDITANDO REGISTRO #{item['id']}"; lbl_titulo_os.color = "orange"
            btn_salvar_os.text = "SALVAR ALTERAÇÕES"; btn_cancelar_edicao.visible = True; ir_para_nova(None)

        def acao_gerar(e):
            btn_gerar.text = "PROCESSANDO..."; btn_gerar.disabled = True; page.update()
            nome_pdf = usuario_atual['nome']
            if dd_filtro_func.visible and dd_filtro_func.value != "todos":
                for opt in dd_filtro_func.options:
                    if opt.key == dd_filtro_func.value: nome_pdf = opt.text
            
            url_pdf = gerar_pdf_nuvem(dados_atuais, f"{txt_dt_ini.value} a {txt_dt_fim.value}", nome_pdf)
            url_word = gerar_word_nuvem(dados_atuais, f"{txt_dt_ini.value} a {txt_dt_fim.value}", nome_pdf)
            
            botoes = []
            if url_pdf:
                link_zap = f"https://wa.me/?text={urllib.parse.quote(f'Olá, segue o relatório: {url_pdf}')}"
                botoes.append(ft.FilledButton("ENVIAR WHATSAPP", url=link_zap, style=ft.ButtonStyle(bgcolor="green"), width=150))
                botoes.append(ft.FilledButton("ABRIR PDF", url=url_pdf, width=150))
            if url_word:
                botoes.append(ft.FilledButton("BAIXAR RELATÓRIO WORD", url=url_word, width=300, style=ft.ButtonStyle(bgcolor="orange")))

            if botoes:
                linha_botoes_pdf.controls = botoes
                linha_botoes_pdf.visible = True
                txt_feedback_pdf.value = "Relatórios prontos! Escolha uma opção:"
                btn_gerar.visible = False 
                
            btn_gerar.text = "GERAR RELATÓRIO"
            btn_gerar.disabled = False
            page.update()

        btn_gerar.on_click = acao_gerar
        
        container_historico = ft.Column([
            ft.Text("FILTRAR PERÍODO", size=14, weight="bold"),
            ft.Row([
                ft.OutlinedButton("HOJE", on_click=lambda e: (setattr(txt_dt_ini, 'value', str(datetime.date.today())), page.update())),
                ft.OutlinedButton("ESTE MÊS", on_click=lambda e: (setattr(txt_dt_ini, 'value', str(datetime.date.today().replace(day=1))), page.update())),
                ft.OutlinedButton("MÊS PASSADO", on_click=lambda e: (setattr(txt_dt_ini, 'value', str((datetime.date.today().replace(day=1) - timedelta(days=1)).replace(day=1))), page.update())),
            ], alignment="center", wrap=True),
            ft.Row([txt_dt_ini, txt_dt_fim]),
            dd_filtro_func, ft.FilledButton("BUSCAR REGISTROS", on_click=buscar, width=300),
            txt_feedback_pdf, btn_gerar, linha_botoes_pdf, lista_cards
        ])

        area_principal = ft.Container(content=container_nova_os)
        def ir_para_nova(e): area_principal.content = container_nova_os; page.update()
        def ir_para_hist(e):
            area_principal.content = container_historico
            setor = usuario_atual['setor'].lower()
            if any(c in setor for c in ['admin', 'dono', 'gerente', 'chefe']):
                dd_filtro_func.visible = True
                if not dd_filtro_func.options:
                    dd_filtro_func.options.append(ft.dropdown.Option(text="TODOS", key="todos"))
                    for u in supabase.table("usuarios").select("id, nome").execute().data:
                        dd_filtro_func.options.append(ft.dropdown.Option(text=u['nome'], key=str(u['id'])))
                    dd_filtro_func.value = "todos"
            page.update()

        page.add(ft.Container(bgcolor="blue", padding=10, content=ft.Row([
            ft.Column([ft.Text(f"Olá, {usuario_atual['nome']}", color="white", weight="bold"), ft.Text(f"Setor: {usuario_atual['setor']}", color="white", size=10)]),
            ft.FilledButton("SAIR", on_click=lambda e: tela_login(), style=ft.ButtonStyle(bgcolor="red"))
        ], alignment="spaceBetween")))
        page.add(ft.Row([ft.FilledButton("NOVA OS", expand=True, on_click=ir_para_nova), ft.FilledButton("HISTORICO", expand=True, on_click=ir_para_hist)]))
        page.add(area_principal)

    tela_login()

ft.app(target=main, view=ft.AppView.WEB_BROWSER, port=int(os.environ.get("PORT", 8550)), host="0.0.0.0")